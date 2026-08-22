"""Source-loss prevention and WebP cover regression tests."""

from __future__ import annotations

import json
import sys
import tempfile
import unittest
import uuid
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage


BOOK_GENERATOR = Path(__file__).resolve().parents[1]
REPO_ROOT = BOOK_GENERATOR.parents[1]
sys.path.insert(0, str(BOOK_GENERATOR))

import run_generation_worker as worker  # noqa: E402
from generate_book import main as generate_book_main  # noqa: E402
from greyveil.importer import block_text, import_docx, parse_manuscript, validate_cover_image  # noqa: E402
from greyveil.loader import load_book  # noqa: E402
from greyveil.qa import validate_generated_outputs  # noqa: E402


class SourceFidelityTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory(prefix="greyveil-source-fidelity-")
        self.temp_path = Path(self.temp_dir.name)
        self.reference_cover = REPO_ROOT / "assets" / "books" / "the-last-shift" / "cover" / "front-cover-print.png"

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_paragraph_only_docx_preserves_unicode_punctuation_and_order(self) -> None:
        document = base_document()
        expected = ["First — unchanged café.", "Second “quoted” paragraph!", "Third line…"]
        for text in expected:
            document.add_paragraph(text)
        parsed = parse_manuscript(document)
        self.assertEqual(parsed["source_errors"], [])
        actual = [block_text(block) for block in parsed["units"][0]["elements"] if block["type"] == "paragraph"]
        self.assertEqual(actual, expected)

    def test_empty_table_is_ignored_without_reordering_surrounding_paragraphs(self) -> None:
        document = Document()
        document.core_properties.title = "Fidelity Fixture"
        document.core_properties.author = "A. Writer"
        document.styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
        document.add_paragraph("Fidelity Fixture", style="Title")
        document.add_paragraph("A. Writer", style="Author")
        document.add_paragraph("Paragraph A")
        empty_table = document.add_table(rows=1, cols=2)
        empty_table.cell(0, 0).text = " \t "
        document.add_paragraph("Paragraph B")
        document.add_paragraph("Heading 1", style="Heading 1")
        document.add_paragraph("Paragraph C")
        parsed = parse_manuscript(document)
        self.assertEqual(parsed["source_errors"], [])
        actual = [
            *[block_text(block) for block in parsed["opening"] if block_text(block).startswith("Paragraph")],
            parsed["units"][0]["title"],
            *[block_text(block) for block in parsed["units"][0]["elements"] if block["type"] == "paragraph"],
        ]
        self.assertEqual(actual, ["Paragraph A", "Paragraph B", "Heading 1", "Paragraph C"])

    def test_supported_word_metadata_unused_numbering_and_plain_url_pass(self) -> None:
        header_art = self.temp_path / "header-art.png"
        PILImage.new("RGB", (12, 12), (20, 80, 140)).save(header_art, "PNG")
        document = base_document()
        paragraph = document.add_paragraph("https://example.com — plain URL text remains supported. ")
        paragraph.add_run("Bold text").bold = True
        paragraph.add_run(" and italic text").italic = True
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "41")
        bookmark_start.set(qn("w:name"), "SupportedBookmark")
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "41")
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)
        proof_start = OxmlElement("w:proofErr")
        proof_start.set(qn("w:type"), "spellStart")
        proof_end = OxmlElement("w:proofErr")
        proof_end.set(qn("w:type"), "spellEnd")
        paragraph._p.insert(1, proof_start)
        paragraph._p.append(proof_end)
        document.sections[0].header.paragraphs[0].add_run().add_picture(str(header_art))
        document.sections[0].footer.paragraphs[0].text = "Footer intentionally outside the body contract."
        manuscript = self.temp_path / "supported-metadata.docx"
        document.save(manuscript)
        reopened = Document(manuscript)

        numbering = reopened.part.numbering_part.element
        self.assertTrue(any(node.tag.rsplit("}", 1)[-1] == "abstractNum" for node in numbering))
        parsed = parse_manuscript(reopened)
        self.assertEqual(parsed["source_errors"], [])
        self.assertTrue(any("https://example.com" in block_text(block) for block in parsed["units"][0]["elements"]))

    def test_intentionally_empty_structural_unit_remains_valid(self) -> None:
        parsed = parse_manuscript(base_document())
        self.assertEqual(parsed["source_errors"], [])
        self.assertEqual(len(parsed["units"]), 1)
        self.assertEqual(parsed["units"][0]["elements"], [])

    def test_nonempty_tables_fail_before_normalized_source_is_created(self) -> None:
        fixtures = {}

        simple = base_document()
        simple.add_paragraph("Before table.")
        table = simple.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Meaningful cell"
        simple.add_paragraph("After table.")
        fixtures["simple"] = simple

        multi = base_document()
        cell = multi.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "First cell paragraph"
        cell.add_paragraph("Second cell paragraph")
        fixtures["multi-paragraph"] = multi

        merged = base_document()
        merged_table = merged.add_table(rows=2, cols=2)
        merged_table.cell(0, 0).merge(merged_table.cell(0, 1)).text = "Merged content"
        fixtures["merged"] = merged

        nested = base_document()
        outer_cell = nested.add_table(rows=1, cols=1).cell(0, 0)
        outer_cell.add_table(rows=1, cols=1).cell(0, 0).text = "Nested content"
        fixtures["nested"] = nested

        for label, document in fixtures.items():
            with self.subTest(label=label):
                manuscript = self.temp_path / f"{label}.docx"
                document.save(manuscript)
                workspace = self.temp_path / f"workspace-{label}"
                result = import_docx(
                    REPO_ROOT,
                    manuscript,
                    cover_path=self.reference_cover,
                    metadata_overrides={"title": f"Table {label}", "author": "A. Writer"},
                    workspace_root=workspace,
                    approved_slug=f"table-{label}",
                )
                self.assertEqual(result.status, "failed")
                self.assertEqual(result.error, "Manuscript contains a Word table that cannot yet be published safely.")
                self.assertFalse((workspace / "assets" / "books" / f"table-{label}").exists())

    def test_meaningful_media_lists_hyperlinks_and_equations_are_not_silently_omitted(self) -> None:
        image_path = self.temp_path / "inline.png"
        PILImage.new("RGB", (12, 12), (20, 80, 140)).save(image_path, "PNG")

        inline_media = base_document()
        inline_media.add_picture(str(image_path))

        floating_media = base_document()
        floating_shape = floating_media.add_paragraph().add_run().add_picture(str(image_path))
        floating_shape._inline.tag = qn("wp:anchor")

        vml_media = base_document()
        vml_media.add_paragraph()._p.append(OxmlElement("w:pict"))

        embedded_object = base_document()
        embedded_object.add_paragraph()._p.append(OxmlElement("w:object"))

        textbox = base_document()
        textbox_content = OxmlElement("w:txbxContent")
        textbox_paragraph = OxmlElement("w:p")
        textbox_run = OxmlElement("w:r")
        textbox_text = OxmlElement("w:t")
        textbox_text.text = "Meaningful text box content"
        textbox_run.append(textbox_text)
        textbox_paragraph.append(textbox_run)
        textbox_content.append(textbox_paragraph)
        textbox.add_paragraph()._p.append(textbox_content)

        list_document = base_document()
        list_document.add_paragraph("A meaningful bullet", style="List Bullet")

        hyperlink = base_document()
        paragraph = hyperlink.add_paragraph("Read ")
        hyperlink_element = OxmlElement("w:hyperlink")
        hyperlink_element.set(qn("r:id"), "rId999")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "the linked source"
        run.append(text)
        hyperlink_element.append(run)
        paragraph._p.append(hyperlink_element)

        equation = base_document()
        equation.add_paragraph()._p.append(OxmlElement("m:oMath"))

        cases = {
            "inline-media": (inline_media, "embedded image or drawing"),
            "floating-media": (floating_media, "embedded image or drawing"),
            "vml-media": (vml_media, "embedded image or drawing"),
            "embedded-object": (embedded_object, "embedded object"),
            "textbox": (textbox, "text box"),
            "list": (list_document, "Word list"),
            "hyperlink": (hyperlink, "hyperlink"),
            "equation": (equation, "equation"),
        }
        for label, (document, message) in cases.items():
            with self.subTest(label=label):
                errors = parse_manuscript(document)["source_errors"]
                self.assertTrue(any(message in error for error in errors), errors)

    def test_other_unsupported_body_content_reports_human_readable_categories(self) -> None:
        fixtures = {}

        notes = base_document()
        notes.add_paragraph()._p.append(OxmlElement("w:footnoteReference"))
        fixtures["notes"] = (notes, "footnotes or endnotes")

        field = base_document()
        instruction = OxmlElement("w:instrText")
        instruction.text = "PAGE"
        field.add_paragraph()._p.append(instruction)
        fixtures["field"] = (field, "dynamic Word field")

        revision = base_document()
        inserted = OxmlElement("w:ins")
        inserted_text = OxmlElement("w:t")
        inserted_text.text = "Tracked insertion"
        inserted.append(inserted_text)
        revision.add_paragraph()._p.append(inserted)
        fixtures["revision"] = (revision, "tracked changes")

        control = base_document()
        content_control = OxmlElement("w:sdt")
        content = OxmlElement("w:sdtContent")
        content_text = OxmlElement("w:t")
        content_text.text = "Controlled body content"
        content.append(content_text)
        content_control.append(content)
        control.element.body.insert(-1, content_control)
        fixtures["control"] = (control, "content control")

        for label, (document, message) in fixtures.items():
            with self.subTest(label=label):
                errors = parse_manuscript(document)["source_errors"]
                self.assertTrue(any(message in error for error in errors), errors)

    def test_valid_webp_cover_uses_lossless_print_derivative_and_original_candidate(self) -> None:
        manuscript = self.temp_path / "webp-regeneration.docx"
        create_fixture(manuscript)
        webp_cover = self.temp_path / "new-cover.webp"
        PILImage.new("RGBA", (96, 144), (31, 73, 119, 211)).save(webp_cover, "WEBP", lossless=True)
        PILImage.new("RGB", (20, 30), (220, 10, 10)).save(self.temp_path / "cover.png", "PNG")
        canonical_cover_bytes = self.reference_cover.read_bytes()
        job_id = str(uuid.uuid4())
        candidate_paths = []

        with patch.object(worker, "storage_upload") as mocked_upload:
            for attempt in range(2):
                with worker.isolated_job_workspace(job_id, self.temp_path) as workspace:
                    self.assertFalse((workspace.candidates / "stale-qa.json").exists())
                    result = import_docx(
                        REPO_ROOT,
                        manuscript,
                        design_from="the-last-shift",
                        cover_path=webp_cover,
                        metadata_overrides={"title": "Fidelity Fixture", "author": "A. Writer"},
                        workspace_root=workspace.repository,
                        approved_slug="the-last-shift",
                    )
                    self.assertEqual(result.status, "imported", result.error or result.warnings)
                    book_json = json.loads((result.book_path / "book.json").read_text(encoding="utf-8"))
                    self.assertEqual(book_json["cover"]["web"], "cover/front-cover.webp")
                    self.assertEqual(book_json["cover"]["source"], "cover/front-cover.webp")
                    self.assertEqual(book_json["cover"]["print"], "cover/front-cover-print.png")
                    self.assertEqual(
                        (result.book_path / "cover" / "front-cover.webp").read_bytes(),
                        webp_cover.read_bytes(),
                    )
                    with PILImage.open(result.book_path / book_json["cover"]["print"]) as print_cover:
                        self.assertEqual(print_cover.format, "PNG")
                        self.assertEqual(print_cover.size, (96, 144))
                        self.assertEqual(print_cover.mode, "RGBA")
                        self.assertEqual(print_cover.getpixel((0, 0))[3], 211)

                    self.assertEqual(
                        generate_book_main([
                            "the-last-shift", "--all",
                            "--repo-root", str(workspace.repository),
                            "--output-dir", str(workspace.outputs),
                        ]),
                        0,
                    )
                    model = load_book(workspace.repository, "the-last-shift")
                    files = {
                        "pdf": workspace.outputs / "pdf" / "the-last-shift.pdf",
                        "epub": workspace.outputs / "epub" / "the-last-shift.epub",
                        "docx": workspace.outputs / "docx" / "the-last-shift-print-editable.docx",
                    }
                    report = validate_generated_outputs(model, files)
                    self.assertTrue(report.ok, report.errors)
                    with zipfile.ZipFile(files["epub"]) as epub:
                        self.assertIn("EPUB/images/cover.png", epub.namelist())

                    candidate_cover = worker.cover_file(workspace.repository, "the-last-shift")
                    self.assertEqual(candidate_cover.suffix, ".webp")
                    candidate = worker.upload_candidate(
                        {"id": job_id}, "cover", candidate_cover, worker.cover_mime(candidate_cover)
                    )
                    candidate_paths.append(candidate["path"])
                    self.assertEqual(candidate["mime_type"], "image/webp")
                    self.assertTrue(candidate["path"].endswith("/cover/front-cover.webp"))
                    if attempt == 0:
                        # Simulate a later failure after candidate upload. The next
                        # attempt must not inherit this stale local QA marker.
                        (workspace.candidates / "stale-qa.json").write_text("not ready", encoding="utf-8")
                self.assertFalse(worker.job_workspace_path(job_id, self.temp_path).exists())

        self.assertEqual(candidate_paths[0], candidate_paths[1])
        self.assertEqual(mocked_upload.call_count, 2)
        self.assertTrue(all(call.kwargs == {"upsert": True} for call in mocked_upload.call_args_list))

        self.assertEqual(self.reference_cover.read_bytes(), canonical_cover_bytes)

    def test_corrupt_or_mislabeled_webp_is_rejected_but_png_and_jpeg_remain_valid(self) -> None:
        malformed_manuscript = self.temp_path / "malformed.docx"
        malformed_manuscript.write_text("not a DOCX package", encoding="utf-8")
        malformed_result = import_docx(REPO_ROOT, malformed_manuscript)
        self.assertEqual(malformed_result.status, "failed")
        self.assertEqual(malformed_result.error, "Manuscript is not a valid DOCX file.")
        self.assertNotIn(str(self.temp_path), malformed_result.error)

        corrupt = self.temp_path / "corrupt.webp"
        corrupt.write_text("not an image", encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "valid PNG, JPEG, or WebP"):
            validate_cover_image(corrupt)
        manuscript = self.temp_path / "corrupt-cover.docx"
        create_fixture(manuscript)
        result = import_docx(
            REPO_ROOT,
            manuscript,
            cover_path=corrupt,
            metadata_overrides={"title": "Fidelity Fixture", "author": "A. Writer"},
            workspace_root=self.temp_path / "corrupt-workspace",
            approved_slug="corrupt-cover",
        )
        self.assertEqual(result.status, "failed")
        self.assertEqual(result.error, "Cover must be a valid PNG, JPEG, or WebP image.")
        self.assertNotIn(str(self.temp_path), result.error)

        mislabeled = self.temp_path / "mislabeled.webp"
        PILImage.new("RGB", (8, 8), (1, 2, 3)).save(mislabeled, "PNG")
        with self.assertRaisesRegex(ValueError, "single-frame PNG, JPEG, or WebP"):
            validate_cover_image(mislabeled)

        valid_webp = self.temp_path / "valid-before-truncation.webp"
        PILImage.new("RGB", (18, 24), (7, 8, 9)).save(valid_webp, "WEBP", lossless=True)
        truncated = self.temp_path / "truncated.webp"
        truncated.write_bytes(valid_webp.read_bytes()[:-12])
        with self.assertRaisesRegex(ValueError, "valid PNG, JPEG, or WebP"):
            validate_cover_image(truncated)

        animated = self.temp_path / "animated.webp"
        first_frame = PILImage.new("RGBA", (18, 24), (10, 20, 30, 255))
        second_frame = PILImage.new("RGBA", (18, 24), (30, 20, 10, 128))
        first_frame.save(
            animated,
            "WEBP",
            save_all=True,
            append_images=[second_frame],
            duration=[100, 100],
            loop=0,
            lossless=True,
        )
        with self.assertRaisesRegex(ValueError, "single-frame PNG, JPEG, or WebP"):
            validate_cover_image(animated)

        resource_limited = self.temp_path / "resource-limited.webp"
        PILImage.new("RGB", (10, 10), (11, 12, 13)).save(resource_limited, "WEBP", lossless=True)
        with patch.object(PILImage, "MAX_IMAGE_PIXELS", 16):
            with self.assertRaisesRegex(ValueError, "valid PNG, JPEG, or WebP"):
                validate_cover_image(resource_limited)

        for suffix, image_format in ((".png", "PNG"), (".jpg", "JPEG")):
            path = self.temp_path / f"valid{suffix}"
            PILImage.new("RGB", (8, 8), (4, 5, 6)).save(path, image_format)
            validate_cover_image(path)

    def test_png_and_jpeg_cover_generation_regression(self) -> None:
        manuscript = self.temp_path / "raster-regression.docx"
        create_fixture(manuscript)

        for suffix, image_format, expected_mime in (
            (".png", "PNG", "image/png"),
            (".jpg", "JPEG", "image/jpeg"),
        ):
            with self.subTest(image_format=image_format):
                slug = f"{image_format.casefold()}-fidelity"
                cover = self.temp_path / f"generation-cover{suffix}"
                PILImage.new("RGB", (72, 108), (31, 73, 119)).save(cover, image_format)
                job_id = str(uuid.uuid4())
                with worker.isolated_job_workspace(job_id, self.temp_path) as workspace:
                    result = import_docx(
                        REPO_ROOT,
                        manuscript,
                        cover_path=cover,
                        metadata_overrides={"title": "Fidelity Fixture", "author": "A. Writer"},
                        workspace_root=workspace.repository,
                        approved_slug=slug,
                    )
                    self.assertEqual(result.status, "imported", result.error or result.warnings)
                    book_json = json.loads((result.book_path / "book.json").read_text(encoding="utf-8"))
                    expected_cover = f"cover/front-cover{suffix}"
                    self.assertEqual(book_json["cover"]["web"], expected_cover)
                    self.assertEqual(book_json["cover"]["print"], expected_cover)
                    self.assertEqual(book_json["cover"]["source"], expected_cover)
                    self.assertFalse((result.book_path / "cover" / "front-cover-print.png").exists())

                    self.assertEqual(
                        generate_book_main([
                            slug, "--all",
                            "--repo-root", str(workspace.repository),
                            "--output-dir", str(workspace.outputs),
                        ]),
                        0,
                    )
                    model = load_book(workspace.repository, slug)
                    files = {
                        "pdf": workspace.outputs / "pdf" / f"{slug}.pdf",
                        "epub": workspace.outputs / "epub" / f"{slug}.epub",
                        "docx": workspace.outputs / "docx" / f"{slug}-print-editable.docx",
                    }
                    report = validate_generated_outputs(model, files)
                    self.assertTrue(report.ok, report.errors)
                    candidate_cover = worker.cover_file(workspace.repository, slug)
                    self.assertEqual(candidate_cover.suffix, suffix)
                    self.assertEqual(worker.cover_mime(candidate_cover), expected_mime)


def base_document() -> Document:
    document = Document()
    document.core_properties.title = "Fidelity Fixture"
    document.core_properties.author = "A. Writer"
    try:
        document.styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        pass
    document.add_paragraph("Fidelity Fixture", style="Title")
    document.add_paragraph("A. Writer", style="Author")
    document.add_paragraph("Chapter 1", style="Heading 1")
    return document


def create_fixture(path: Path) -> None:
    document = base_document()
    document.add_paragraph("The WebP replacement cover belongs only to this candidate.")
    document.add_paragraph("Epilogue", style="Heading 1")
    document.add_paragraph("All generated formats retain their structural QA contract.")
    document.save(path)


if __name__ == "__main__":
    unittest.main()
