"""High-value integration test for the DOCX manuscript production pipeline."""

from __future__ import annotations

import json
import shutil
import sys
import tempfile
import unittest
from argparse import Namespace
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


BOOK_GENERATOR = Path(__file__).resolve().parents[1]
REPO_ROOT = BOOK_GENERATOR.parents[1]
sys.path.insert(0, str(BOOK_GENERATOR))

from greyveil.importer import import_docx  # noqa: E402
from greyveil.loader import load_book  # noqa: E402
from greyveil.qa import validate_generated_outputs  # noqa: E402
from import_book import inbox_manuscripts, process_import  # noqa: E402


class ImportPipelineTest(unittest.TestCase):
    slug = "the-compass-within"

    def setUp(self) -> None:
        self.destination = REPO_ROOT / "assets" / "books" / self.slug
        if self.destination.exists():
            self.fail(f"Test destination already exists and will not be overwritten: {self.destination}")
        self.temp_dir = tempfile.TemporaryDirectory(prefix="greyveil-import-fixture-")
        self.temp_path = Path(self.temp_dir.name)
        self.manuscript = self.temp_path / "the-compass-within.docx"
        self.output_dir = self.temp_path / "output"
        create_fixture(self.manuscript)

    def tearDown(self) -> None:
        if self.destination.exists():
            shutil.rmtree(self.destination)
        self.temp_dir.cleanup()

    def test_docx_import_and_existing_generation(self) -> None:
        cover = REPO_ROOT / "assets" / "books" / "the-last-shift" / "cover" / "front-cover-print.png"
        result = process_import(
            REPO_ROOT,
            self.manuscript,
            Namespace(
                cover=cover,
                design_from="the-last-shift",
                generate=True,
                output_dir=self.output_dir,
            ),
        )

        self.assertEqual(result.status, "complete", result.error or "; ".join(result.warnings))
        self.assertEqual(result.slug, self.slug)
        self.assertEqual(result.metadata["title"], "The Compass Within")
        self.assertEqual(result.metadata["author"], "A. Writer")
        self.assertTrue(result.book_path and result.book_path.exists())

        book_json = json.loads((self.destination / "book.json").read_text(encoding="utf-8"))
        self.assertEqual(book_json["slug"], self.slug)
        self.assertEqual([unit["title"] for unit in book_json["units"]], ["The Compass Within", "Chapter 1", "Chapter 2", "Epilogue"])

        opening = json.loads((self.destination / "chapters" / "00-opening.json").read_text(encoding="utf-8"))
        self.assertEqual(opening["openingMode"], "source")
        self.assertEqual("".join(run["text"] for run in opening["elements"][0]["runs"]), "The Compass Within")

        chapter_one = json.loads((self.destination / "chapters" / "01-chapter-1.json").read_text(encoding="utf-8"))
        paragraph = next(element for element in chapter_one["elements"] if element["type"] == "paragraph")
        self.assertEqual("".join(run["text"] for run in paragraph["runs"]), "Plain bold and italic prose.")
        self.assertTrue(any(run.get("bold") for run in paragraph["runs"]))
        self.assertTrue(any(run.get("italic") for run in paragraph["runs"]))

        model = load_book(REPO_ROOT, self.slug)
        self.assertFalse([issue for issue in model.issues if issue.severity == "error"])
        report = validate_generated_outputs(
            model,
            {
                "pdf": self.output_dir / "pdf" / f"{self.slug}.pdf",
                "epub": self.output_dir / "epub" / f"{self.slug}.epub",
                "docx": self.output_dir / "docx" / f"{self.slug}-print-editable.docx",
            },
        )
        self.assertTrue(report.ok, report.errors)

        conflict = import_docx(REPO_ROOT, self.manuscript, cover_path=cover)
        self.assertEqual(conflict.status, "needs_attention")
        self.assertTrue(any("Destination already exists" in warning for warning in conflict.warnings))

        temporary_word_file = self.temp_path / "~$the-compass-within.docx"
        temporary_word_file.write_bytes(b"temporary")
        self.assertEqual(inbox_manuscripts(self.temp_path), [self.manuscript])


def create_fixture(path: Path) -> None:
    document = Document()
    document.core_properties.title = "The Compass Within"
    document.core_properties.subject = "A practical meditation"
    document.core_properties.author = "A. Writer"
    try:
        document.styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        pass
    document.add_paragraph("The Compass Within", style="Title")
    document.add_paragraph("A practical meditation", style="Subtitle")
    document.add_paragraph("A. Writer", style="Author")
    document.add_page_break()
    document.add_paragraph("Chapter 1", style="Heading 1")
    paragraph = document.add_paragraph()
    paragraph.add_run("Plain ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(" prose.")
    document.add_paragraph("Chapter 2", style="Heading 1")
    document.add_paragraph("The second chapter remains exactly as written.")
    document.add_paragraph("Epilogue", style="Heading 1")
    document.add_paragraph("The final reflection closes the manuscript.")
    document.save(path)


if __name__ == "__main__":
    unittest.main()
