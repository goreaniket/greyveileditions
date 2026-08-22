"""Executable retry, workspace, slug-authority, and path-isolation checks."""

from __future__ import annotations

import sys
import tempfile
import unittest
import uuid
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


BOOK_GENERATOR = Path(__file__).resolve().parents[1]
REPO_ROOT = BOOK_GENERATOR.parents[1]
sys.path.insert(0, str(BOOK_GENERATOR))

import run_generation_worker as worker  # noqa: E402
from generate_book import main as generate_book_main  # noqa: E402
from greyveil.importer import import_docx, normalize_approved_slug  # noqa: E402
from greyveil.loader import load_book  # noqa: E402
from greyveil.qa import validate_generated_outputs  # noqa: E402


class WorkerRetrySafetyTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory(prefix="greyveil-worker-retry-")
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_existing_book_regeneration_retries_without_touching_repository_source(self) -> None:
        slug = "the-last-shift"
        job_id = str(uuid.uuid4())
        manuscript = self.temp_path / "replacement.docx"
        cover = REPO_ROOT / "assets" / "books" / slug / "cover" / "front-cover-print.png"
        canonical_book_json = REPO_ROOT / "assets" / "books" / slug / "book.json"
        original_bytes = canonical_book_json.read_bytes()
        create_fixture(manuscript)

        for attempt in range(2):
            with worker.isolated_job_workspace(job_id, self.temp_path) as workspace:
                self.assertEqual(workspace.root, worker.job_workspace_path(job_id, self.temp_path))
                result = import_docx(
                    REPO_ROOT,
                    manuscript,
                    design_from=slug,
                    cover_path=cover,
                    metadata_overrides={"title": "Replacement Edition", "author": "A. Writer"},
                    workspace_root=workspace.repository,
                    approved_slug=slug,
                )
                self.assertEqual(result.status, "imported", result.error or result.warnings)
                self.assertEqual(result.slug, slug)
                self.assertTrue((workspace.repository / "assets" / "books" / slug / "book.json").is_file())
                self.assertEqual(
                    generate_book_main([
                        slug,
                        "--all",
                        "--repo-root", str(workspace.repository),
                        "--output-dir", str(workspace.outputs),
                    ]),
                    0,
                )
                model = load_book(workspace.repository, slug)
                files = {
                    "pdf": workspace.outputs / "pdf" / f"{slug}.pdf",
                    "epub": workspace.outputs / "epub" / f"{slug}.epub",
                    "docx": workspace.outputs / "docx" / f"{slug}-print-editable.docx",
                }
                report = validate_generated_outputs(model, files)
                self.assertTrue(report.ok, report.errors)
                (workspace.root / f"attempt-{attempt}").write_text("temporary", encoding="utf-8")
            self.assertFalse(worker.job_workspace_path(job_id, self.temp_path).exists())

        self.assertEqual(canonical_book_json.read_bytes(), original_bytes)

    def test_workspace_clears_stale_same_job_state_before_retry(self) -> None:
        job_id = str(uuid.uuid4())
        stale_root = worker.job_workspace_path(job_id, self.temp_path)
        stale_root.mkdir(parents=True)
        (stale_root / "stale.txt").write_text("old attempt", encoding="utf-8")
        with worker.isolated_job_workspace(job_id, self.temp_path) as workspace:
            self.assertFalse((workspace.root / "stale.txt").exists())
            self.assertTrue(workspace.inputs.is_dir())
            self.assertTrue(workspace.repository.is_dir())
            self.assertTrue(workspace.outputs.is_dir())
            self.assertTrue(workspace.candidates.is_dir())
        self.assertFalse(stale_root.exists())

    def test_workspace_rejects_non_temp_roots_and_redirected_paths(self) -> None:
        job_id = str(uuid.uuid4())
        other_id = str(uuid.uuid4())
        with self.assertRaises(RuntimeError):
            worker.job_workspace_path(job_id, REPO_ROOT)

        trusted = self.temp_path / "trusted"
        trusted.mkdir()
        outside = self.temp_path / "outside"
        outside.mkdir()
        redirected_base = trusted / "greyveil-generation"
        redirected_base.symlink_to(outside, target_is_directory=True)
        with self.assertRaises(RuntimeError):
            worker.job_workspace_path(job_id, trusted)
        redirected_base.unlink()

        base = trusted / "greyveil-generation"
        base.mkdir()
        other_workspace = base / other_id
        other_workspace.mkdir()
        redirected_job = base / job_id
        redirected_job.symlink_to(other_workspace, target_is_directory=True)
        with self.assertRaises(RuntimeError):
            worker.job_workspace_path(job_id, trusted)
        with self.assertRaises(RuntimeError):
            worker.remove_job_workspace(job_id, trusted)
        self.assertTrue(other_workspace.is_dir())

        for unsafe_id in ("", "../", "/", "C:\\", r"\\server\share"):
            with self.subTest(unsafe_id=unsafe_id), self.assertRaises(RuntimeError):
                worker.job_workspace_path(unsafe_id, trusted)

    def test_input_paths_are_exactly_bound_to_current_job(self) -> None:
        job_id = str(uuid.uuid4())
        other_id = str(uuid.uuid4())
        valid = {
            "id": job_id,
            "manuscript_path": f"jobs/{job_id}/manuscript.docx",
            "cover_path": f"jobs/{job_id}/cover.webp",
        }
        self.assertEqual(worker.validate_job_input_paths(valid), (valid["manuscript_path"], valid["cover_path"]))
        for unsafe in (
            {**valid, "manuscript_path": f"jobs/{other_id}/manuscript.docx"},
            {**valid, "manuscript_path": f"jobs/{job_id}/../manuscript.docx"},
            {**valid, "manuscript_path": f"jobs/{job_id}/%2e%2e/manuscript.docx"},
            {**valid, "manuscript_path": f"JOBS/{job_id}/manuscript.docx"},
            {**valid, "cover_path": f"jobs/{other_id}/cover.webp"},
            {**valid, "cover_path": f"jobs/{job_id}/nested/cover.png"},
            {**valid, "cover_path": f"jobs/{job_id}/cover.png%2f.."},
        ):
            with self.subTest(unsafe=unsafe), self.assertRaises(RuntimeError):
                worker.validate_job_input_paths(unsafe)

    def test_authoritative_slug_uses_book_for_regeneration_and_admin_for_create(self) -> None:
        with patch.object(worker, "request", return_value=[{"slug": "canonical-book"}]) as mocked_request:
            self.assertEqual(worker.authoritative_job_slug({"book_id": 41, "metadata": {"slug": "ignored"}}), "canonical-book")
            mocked_request.assert_called_once_with("/rest/v1/books?id=eq.41&select=slug")
        self.assertEqual(worker.authoritative_job_slug({"book_id": None, "metadata": {"slug": "admin-approved"}}), "admin-approved")
        for unsafe in (
            "", "../escape", r"..\escape", "two--hyphens", "spaces here", "slash/name", r"slash\name",
            "%2e%2e", "%2fescape", "%5cescape", "-leading", "trailing-", r"C:\book", r"\\server\book",
        ):
            with self.subTest(unsafe=unsafe), self.assertRaises(ValueError):
                normalize_approved_slug(unsafe)

    def test_same_job_candidate_retry_upserts_same_path_but_jobs_remain_isolated(self) -> None:
        first_id = str(uuid.uuid4())
        second_id = str(uuid.uuid4())
        artifact = self.temp_path / "book.pdf"
        artifact.write_bytes(b"candidate")
        job = {"id": first_id}
        with patch.object(worker, "storage_upload") as mocked_upload:
            first = worker.upload_candidate(job, "pdf", artifact, "application/pdf")
            second = worker.upload_candidate(job, "pdf", artifact, "application/pdf")
        self.assertEqual(first, second)
        self.assertEqual(first["path"], f"jobs/{first_id}/pdf/book.pdf")
        self.assertEqual(mocked_upload.call_count, 2)
        self.assertTrue(all(call.kwargs == {"upsert": True} for call in mocked_upload.call_args_list))
        self.assertNotEqual(
            worker.candidate_object_path(first_id, "pdf", artifact.name),
            worker.candidate_object_path(second_id, "pdf", artifact.name),
        )
        with self.assertRaises(RuntimeError):
            worker.candidate_object_path(first_id, "pdf", "../book.pdf")
        with self.assertRaises(RuntimeError):
            worker.candidate_object_path(first_id, "pdf", "%2e%2e-book.pdf")


def create_fixture(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Replacement Edition"
    document.core_properties.author = "A. Writer"
    try:
        document.styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        pass
    document.add_paragraph("Replacement Edition", style="Title")
    document.add_paragraph("A. Writer", style="Author")
    document.add_paragraph("Chapter 1", style="Heading 1")
    document.add_paragraph("The replacement manuscript remains exactly as supplied.")
    document.add_paragraph("Epilogue", style="Heading 1")
    document.add_paragraph("The retry finishes cleanly.")
    document.save(path)


if __name__ == "__main__":
    unittest.main()
