"""Conservative DOCX-to-Greyveil source import.

The importer preserves Word text and semantic structure; it does not rewrite
manuscript prose or create a second rendering pipeline.
"""

from __future__ import annotations

import copy
import json
import os
import re
import shutil
import tempfile
import warnings
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image as PILImage
from PIL import UnidentifiedImageError

from .jobs import GenerationJob, GenerationStage, utc_now
from .loader import load_book


IMPORTER_SCHEMA_VERSION = "greyveil.docx-import/v1"
SAFE_DEFAULT_PUBLISHER = "Greyveil Editions"
IMAGE_SUFFIXES = {"image/jpeg": ".jpg", "image/png": ".png", "image/gif": ".gif", "image/webp": ".webp"}
SUPPORTED_COVER_FORMATS = {".jpg": "JPEG", ".jpeg": "JPEG", ".png": "PNG", ".webp": "WEBP"}
UNSUPPORTED_SOURCE_MESSAGES = {
    "table": "Manuscript contains a Word table that cannot yet be published safely.",
    "media": "Manuscript contains an embedded image or drawing that cannot yet be published safely.",
    "textbox": "Manuscript contains a text box that cannot yet be published safely.",
    "hyperlink": "Manuscript contains a hyperlink that cannot yet be published safely.",
    "list": "Manuscript contains a Word list whose numbering cannot yet be published safely.",
    "notes": "Manuscript contains footnotes or endnotes that cannot yet be published safely.",
    "equation": "Manuscript contains an equation that cannot yet be published safely.",
    "object": "Manuscript contains an embedded object that cannot yet be published safely.",
    "field": "Manuscript contains a dynamic Word field that cannot yet be published safely.",
    "revision": "Manuscript contains tracked changes that must be accepted or rejected before publishing.",
    "control": "Manuscript contains a content control that cannot yet be published safely.",
}
STRUCTURE_PATTERNS = (
    (re.compile(r"^chapter\s+(?:\d+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b", re.I), "chapter"),
    (re.compile(r"^(?:part|arc|phase)\s+(?:\d+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b", re.I), "part"),
)
SPECIAL_UNITS = {
    "copyright": "frontmatter",
    "disclaimer": "frontmatter",
    "dedication": "dedication",
    "author's note": "frontmatter",
    "authors note": "frontmatter",
    "prologue": "prologue",
    "introduction": "introduction",
    "epilogue": "epilogue",
    "final reflection": "ending",
    "about the author": "frontmatter",
}


@dataclass
class ImportResult:
    status: str
    job: GenerationJob
    manuscript: Path
    slug: str = ""
    book_path: Path | None = None
    metadata: dict[str, str] = field(default_factory=dict)
    missing_fields: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    error: str = ""


def slugify(value: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "-", value.casefold()).strip("-")
    return normalized or "untitled-book"


def import_docx(
    repo_root: Path,
    manuscript_path: Path,
    *,
    design_from: str = "the-last-shift",
    cover_path: Path | None = None,
    metadata_overrides: dict[str, str] | None = None,
    workspace_root: Path | None = None,
    approved_slug: str | None = None,
) -> ImportResult:
    job = GenerationJob()
    manuscript_path = manuscript_path.resolve()
    job.advance(GenerationStage.IMPORTING)
    if manuscript_path.suffix.casefold() != ".docx" or not manuscript_path.is_file():
        return failed_result(job, manuscript_path, "Input must be an existing .docx manuscript.")

    try:
        document = Document(manuscript_path)
    except Exception:  # noqa: BLE001 - do not expose parser or filesystem details.
        return failed_result(job, manuscript_path, "Manuscript is not a valid DOCX file.")

    job.advance(GenerationStage.NORMALIZING)
    parsed = parse_manuscript(document)
    metadata = detect_metadata(document, parsed["opening"])
    # Founder review happens before generation. Only explicit, non-empty admin
    # corrections replace detected values; manuscript text is never rewritten.
    for key, value in (metadata_overrides or {}).items():
        if isinstance(value, str) and value.strip():
            metadata[key] = value.strip()
    title = metadata.get("title", "")
    try:
        slug = normalize_approved_slug(approved_slug) if approved_slug is not None else (slugify(title) if title else "")
    except ValueError as exc:
        return failed_result(job, manuscript_path, str(exc), metadata=metadata)
    if slug:
        metadata["slug"] = slug
    if parsed["source_errors"]:
        return failed_result(job, manuscript_path, "; ".join(parsed["source_errors"]), slug, metadata)
    try:
        selected_cover = find_cover(manuscript_path, cover_path)
    except ValueError as exc:
        return failed_result(job, manuscript_path, str(exc), slug, metadata)
    missing_fields = []
    if not title:
        missing_fields.append("title")
    if not metadata.get("author"):
        missing_fields.append("author")
    if not selected_cover:
        missing_fields.append("cover")
    if not parsed["units"]:
        missing_fields.append("manuscript structure")
    if missing_fields:
        message = "Needs founder confirmation: " + ", ".join(missing_fields) + "."
        job.add_warning(message)
        return ImportResult(
            status="needs_attention",
            job=job,
            manuscript=manuscript_path,
            slug=slug,
            metadata=metadata,
            missing_fields=missing_fields,
            warnings=list(job.warnings),
        )

    generation_root = (workspace_root or repo_root).resolve()
    destination = generation_root / "assets" / "books" / slug
    if destination.exists():
        return attention_result(job, manuscript_path, slug, metadata, "Destination already exists.")

    reference = repo_root / "assets" / "books" / design_from
    if not (reference / "design-spec.json").is_file() or not (reference / "theme.css").is_file():
        return failed_result(job, manuscript_path, f"Design reference is not usable: {design_from}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    staging: Path | None = Path(tempfile.mkdtemp(prefix=f".{slug}.import-", dir=destination.parent))
    try:
        write_imported_source(staging, parsed, metadata, slug, reference, selected_cover, manuscript_path)
        job.advance(GenerationStage.VALIDATING_SOURCE)
        staging_model = load_book(generation_root, staging.name)
        errors = [issue.message for issue in staging_model.issues if issue.severity == "error"]
        if errors:
            return failed_result(job, manuscript_path, "Imported source did not pass validation.", slug, metadata)
        for issue in staging_model.issues:
            if issue.severity == "warning":
                job.add_warning(issue.message)
        if destination.exists():
            return attention_result(job, manuscript_path, slug, metadata, "Destination already exists.")
        os.replace(staging, destination)
        staging = None
        return ImportResult(
            status="imported",
            job=job,
            manuscript=manuscript_path,
            slug=slug,
            book_path=destination,
            metadata=metadata,
            warnings=list(job.warnings),
        )
    except Exception:  # noqa: BLE001 - never expose library or filesystem details.
        return failed_result(job, manuscript_path, "Could not finalize imported source safely.", slug, metadata)
    finally:
        if staging is not None and staging.exists():
            shutil.rmtree(staging, ignore_errors=True)


def normalize_approved_slug(value: str | None) -> str:
    normalized = str(value or "").strip().casefold()
    if not re.fullmatch(r"[a-z0-9]+(?:-[a-z0-9]+)*", normalized):
        raise ValueError("Approved slug must contain lowercase letters, numbers, and single hyphens only.")
    return normalized


def parse_manuscript(document: Document) -> dict[str, Any]:
    opening: list[dict[str, Any]] = []
    units: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    image_counter = 0
    paragraph_index = 0
    source_errors = audit_source_fidelity(document)

    # document.paragraphs and document.tables are separate collections. Walking
    # inner content keeps body order stable and makes table omission explicit.
    for body_block in document.iter_inner_content():
        if isinstance(body_block, Table):
            continue
        if not isinstance(body_block, Paragraph):
            continue
        paragraph = body_block
        paragraph_index += 1
        text = paragraph.text
        style_name = paragraph.style.name if paragraph.style else ""
        classification = classify_structure(text, style_name)
        if classification and classification[0] != "subheading":
            current = new_unit(classification[0], text, paragraph_index, len(units))
            units.append(current)
            continue

        block_type = "heading" if classification and classification[0] == "subheading" else paragraph_block_type(style_name)
        blocks = blocks_from_paragraph(paragraph, paragraph_index, block_type)
        image_blocks = image_blocks_from_paragraph(paragraph, paragraph_index, image_counter)
        image_counter += len(image_blocks)
        blocks.extend(image_blocks)
        if not blocks:
            continue

        if current is None:
            opening.extend(blocks)
        else:
            current["elements"].extend(blocks)

    if current is None and opening:
        # A title-page-only document is not a usable manuscript; its structure is
        # reported as needs_attention by the caller instead of inventing chapters.
        return {"opening": opening, "units": [], "source_errors": source_errors}
    return {"opening": opening, "units": units, "source_errors": source_errors}


def audit_source_fidelity(document: Document) -> list[str]:
    """Reject meaningful Word structures the normalized/export models cannot preserve."""
    body = document.element.body
    categories: list[str] = []

    if any(xml_has_meaningful_text(table) for table in body.iter(qn("w:tbl"))):
        categories.append("table")

    local_names = {element.tag.rsplit("}", 1)[-1] for element in body.iter()}
    if local_names.intersection({"drawing", "pict"}):
        categories.append("media")
    if "txbxContent" in local_names:
        categories.append("textbox")
    if "hyperlink" in local_names:
        categories.append("hyperlink")
    if local_names.intersection({"footnoteReference", "endnoteReference"}):
        categories.append("notes")
    if local_names.intersection({"oMath", "oMathPara"}):
        categories.append("equation")
    if local_names.intersection({"object", "OLEObject", "altChunk", "control", "subDoc"}):
        categories.append("object")
    if local_names.intersection({"fldSimple", "instrText"}):
        categories.append("field")
    if local_names.intersection({"ins", "del", "moveFrom", "moveTo"}):
        categories.append("revision")
    if any(xml_has_meaningful_text(control) for control in body.iter(qn("w:sdt"))):
        categories.append("control")

    for body_block in document.iter_inner_content():
        if not isinstance(body_block, Paragraph):
            continue
        style_name = body_block.style.name.casefold() if body_block.style else ""
        p_pr = body_block._p.pPr
        if (
            style_name.startswith("list ")
            or (p_pr is not None and p_pr.numPr is not None)
        ):
            categories.append("list")
            break

    return [UNSUPPORTED_SOURCE_MESSAGES[category] for category in dict.fromkeys(categories)]


def xml_has_meaningful_text(element) -> bool:
    return any(
        node.text and node.text.strip()
        for node in element.iter()
        if node.tag.rsplit("}", 1)[-1] in {"t", "delText"}
    )


def classify_structure(text: str, style_name: str) -> tuple[str, str] | None:
    normalized = text.strip()
    if not normalized:
        return None
    style = style_name.casefold()
    if style.startswith("heading 2") or style.startswith("heading 3"):
        return ("subheading", normalized)
    for label, kind in SPECIAL_UNITS.items():
        if normalized.casefold() == label:
            return (kind, normalized)
    for pattern, kind in STRUCTURE_PATTERNS:
        if pattern.match(normalized):
            return (kind, normalized)
    if style.startswith("heading 1"):
        return ("chapter", normalized)
    return None


def paragraph_block_type(style_name: str) -> str:
    style = style_name.casefold()
    if "quote" in style or "epigraph" in style:
        return "quote"
    return "paragraph"


def blocks_from_paragraph(paragraph, paragraph_index: int, block_type: str) -> list[dict[str, Any]]:
    text = paragraph.text
    has_page_break = paragraph.paragraph_format.page_break_before or any(run_has_page_break(run) for run in paragraph.runs)
    if not text:
        blocks = []
        if has_page_break:
            blocks.append({"type": "section-break", "sourceParagraph": paragraph_index, "sourcePageBreak": True})
        blocks.append({"type": "space", "sourceParagraph": paragraph_index})
        return blocks
    blocks: list[dict[str, Any]] = []
    if has_page_break:
        blocks.append({"type": "section-break", "sourceParagraph": paragraph_index, "sourcePageBreak": True})
    runs = [
        {key: value for key, value in {"text": run.text, "bold": run.bold, "italic": run.italic}.items() if value not in (None, "")}
        for run in paragraph.runs
        if run.text
    ]
    block: dict[str, Any] = {"type": block_type, "runs": runs or [{"text": text}], "sourceParagraph": paragraph_index}
    role = opening_role(paragraph.style.name if paragraph.style else "", text)
    if role:
        block["role"] = role
    if block_type == "heading":
        block["level"] = heading_level(paragraph.style.name if paragraph.style else "")
    blocks.append(block)
    return blocks


def run_has_page_break(run) -> bool:
    return any(
        element.tag.endswith("}br") and element.get(qn("w:type")) == "page"
        for element in run._element.iter()
    )


def image_blocks_from_paragraph(paragraph, paragraph_index: int, start_index: int) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    try:
        blips = paragraph._p.xpath(".//a:blip")
    except Exception:  # pragma: no cover - depends on nonstandard Word XML.
        return blocks
    for offset, blip in enumerate(blips, start=1):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            continue
        image_part = paragraph.part.related_parts.get(rel_id)
        if not image_part:
            continue
        blocks.append(
            {
                "type": "image",
                "sourceParagraph": paragraph_index,
                "_imageBlob": image_part.blob,
                "_imageSuffix": IMAGE_SUFFIXES.get(image_part.content_type, ".bin"),
                "_imageIndex": start_index + offset,
            }
        )
    return blocks


def opening_role(style_name: str, text: str) -> str:
    style = style_name.casefold()
    if style == "title" or "book title" in style:
        return "title"
    if style == "subtitle" or "book subtitle" in style:
        return "subtitle-line"
    if "author" in style:
        return "author"
    if "collection" in style:
        return "collection-line"
    if "series" in style:
        return "series-line"
    if "publisher" in style or "imprint" in style:
        return "publisher"
    if "book number" in style:
        return "book-number"
    if re.fullmatch(r"book\s+[ivxlcdm\d.]+", text.strip(), re.I):
        return "book-number"
    return ""


def heading_level(style_name: str) -> int:
    match = re.search(r"heading\s+([123])", style_name, flags=re.I)
    return int(match.group(1)) if match else 2


def new_unit(kind: str, title: str, paragraph_index: int, index: int) -> dict[str, Any]:
    identifier = slugify(title) or f"unit-{index + 1:02d}"
    return {
        "id": identifier,
        "kind": kind,
        "label": title,
        "title": title,
        "shortTitle": title,
        "index": index,
        "sourceHeadingParagraphs": [paragraph_index],
        "sourceRange": [paragraph_index, paragraph_index],
        "elements": [],
    }


def detect_metadata(document: Document, opening: list[dict[str, Any]]) -> dict[str, str]:
    properties = document.core_properties
    opening_values: dict[str, str] = {}
    for block in opening:
        role = str(block.get("role", ""))
        text = block_text(block)
        if role == "title" and "title" not in opening_values:
            opening_values["title"] = text
        elif role == "subtitle-line" and "subtitle" not in opening_values:
            opening_values["subtitle"] = text
        elif role == "author" and "author" not in opening_values:
            opening_values["author"] = text
        elif role == "collection-line" and "collection" not in opening_values:
            opening_values["collection"] = text
        elif role == "series-line" and "series" not in opening_values:
            opening_values["series"] = text
        elif role == "book-number" and "bookNumber" not in opening_values:
            opening_values["bookNumber"] = text
        elif role == "publisher" and "publisher" not in opening_values:
            opening_values["publisher"] = text
    core = {
        "title": string_value(properties.title),
        "subtitle": string_value(properties.subject),
        "author": string_value(properties.author),
        "language": string_value(getattr(properties, "language", "")),
    }
    metadata = {key: value for key, value in {**opening_values, **{key: value for key, value in core.items() if value}}.items() if value}
    metadata.setdefault("publisher", SAFE_DEFAULT_PUBLISHER)
    return metadata


def string_value(value: Any) -> str:
    return value.strip() if isinstance(value, str) else ""


def block_text(block: dict[str, Any]) -> str:
    return "".join(str(run.get("text", "")) for run in block.get("runs", []) if isinstance(run, dict)).strip()


def find_cover(manuscript: Path, explicit_cover: Path | None) -> Path | None:
    if explicit_cover:
        candidate = explicit_cover.resolve()
        if not candidate.is_file():
            return None
        validate_cover_image(candidate)
        return candidate
    names = (
        f"{manuscript.stem}-cover.png",
        f"{manuscript.stem}-cover.jpg",
        f"{manuscript.stem}-cover.jpeg",
        f"{manuscript.stem}-cover.webp",
        "cover.png",
        "cover.jpg",
        "cover.jpeg",
        "cover.webp",
    )
    available = {candidate.name.casefold(): candidate for candidate in manuscript.parent.iterdir() if candidate.is_file()}
    for name in names:
        candidate = available.get(name.casefold())
        if candidate:
            validate_cover_image(candidate)
            return candidate
    return None


def validate_cover_image(path: Path) -> None:
    expected_format = SUPPORTED_COVER_FORMATS.get(path.suffix.casefold())
    if not expected_format:
        raise ValueError("Cover must use PNG, JPEG, or WebP format.")
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", PILImage.DecompressionBombWarning)
            with PILImage.open(path) as image:
                image.verify()
            with PILImage.open(path) as image:
                actual_format = image.format
                dimensions = image.size
                frame_count = getattr(image, "n_frames", 1)
                image.load()
    except (
        PILImage.DecompressionBombError,
        PILImage.DecompressionBombWarning,
        UnidentifiedImageError,
        OSError,
        ValueError,
    ) as exc:
        raise ValueError("Cover must be a valid PNG, JPEG, or WebP image.") from exc
    if actual_format != expected_format or min(dimensions) < 1 or frame_count != 1:
        raise ValueError("Cover must be a valid single-frame PNG, JPEG, or WebP image.")


def write_imported_source(
    target: Path,
    parsed: dict[str, Any],
    metadata: dict[str, str],
    slug: str,
    reference: Path,
    cover_source: Path,
    manuscript_path: Path,
) -> None:
    # The caller reserves this uniquely named staging directory before writing.
    target.mkdir(parents=True, exist_ok=True)
    (target / "chapters").mkdir()
    (target / "cover").mkdir()
    shutil.copy2(reference / "theme.css", target / "theme.css")
    write_json(target / "design-spec.json", imported_design(reference / "design-spec.json", metadata, slug, cover_source))

    cover_roles = materialize_cover_assets(target / "cover", cover_source)
    # Relative cover roles validate in the staging directory and remain valid
    # once that directory is atomically renamed to the final slug.
    units = materialize_units(parsed, target, metadata, slug)
    book_json: dict[str, Any] = {
        "id": slug,
        "slug": slug,
        "title": metadata["title"],
        "author": metadata["author"],
        "publisher": metadata.get("publisher", SAFE_DEFAULT_PUBLISHER),
        "designSpecFile": "design-spec.json",
        "themeStylesheet": "theme.css",
        "cover": {
            **cover_roles,
            "alt": f"{metadata['title']} book cover",
        },
        "units": [unit_summary(unit) for unit in units],
    }
    for key in ("subtitle", "collection", "series", "bookNumber", "language", "copyright"):
        if metadata.get(key):
            book_json[key] = metadata[key]
    write_json(target / "book.json", book_json)
    write_json(
        target / "import-manifest.json",
        {
            "schemaVersion": IMPORTER_SCHEMA_VERSION,
            "sourceManuscript": manuscript_path.name,
            "importedAt": utc_now(),
            "detectedMetadata": metadata,
            "designReference": reference.name,
            "coverSource": cover_source.name,
            "warnings": [],
        },
    )


def materialize_cover_assets(cover_dir: Path, cover_source: Path) -> dict[str, str]:
    suffix = cover_source.suffix.casefold()
    web_name = f"front-cover{suffix}"
    shutil.copy2(cover_source, cover_dir / web_name)
    web_path = f"cover/{web_name}"
    if suffix != ".webp":
        return {"web": web_path, "print": web_path, "source": web_path}

    print_name = "front-cover-print.png"
    try:
        with PILImage.open(cover_source) as image:
            image.load()
            converted = image.convert("RGBA" if "A" in image.getbands() else "RGB")
            save_options: dict[str, Any] = {"format": "PNG", "compress_level": 9}
            if image.info.get("icc_profile"):
                save_options["icc_profile"] = image.info["icc_profile"]
            converted.save(cover_dir / print_name, **save_options)
    except Exception as exc:  # noqa: BLE001 - never expose decoder or filesystem details.
        raise ValueError("Cover image could not be prepared safely.") from exc
    return {"web": web_path, "print": f"cover/{print_name}", "source": web_path}


def materialize_units(
    parsed: dict[str, Any], target: Path, metadata: dict[str, str], slug: str
) -> list[dict[str, Any]]:
    units: list[dict[str, Any]] = []
    if any(block.get("role") == "title" for block in parsed["opening"]):
        opening = {
            "id": "opening",
            "kind": "opening",
            "label": "Opening",
            "title": metadata["title"],
            "shortTitle": "Opening",
            "index": 0,
            "openingMode": "source",
            "elements": parsed["opening"],
        }
        units.append(opening)
    units.extend(copy.deepcopy(parsed["units"]))
    seen_ids: set[str] = set()
    for index, unit in enumerate(units):
        base = unit["id"]
        suffix = 2
        while unit["id"] in seen_ids:
            unit["id"] = f"{base}-{suffix}"
            suffix += 1
        seen_ids.add(unit["id"])
        unit["index"] = index
        unit["file"] = f"chapters/{index:02d}-{unit['id']}.json"
        unit["sourceRange"] = source_range(unit)
        materialize_inline_images(unit["elements"], target)
        write_json(target / unit["file"], unit)
    return units


def materialize_inline_images(elements: list[dict[str, Any]], target: Path) -> None:
    image_dir = target / "images"
    for block in elements:
        blob = block.pop("_imageBlob", None)
        suffix = block.pop("_imageSuffix", "")
        index = block.pop("_imageIndex", None)
        if blob is None or index is None:
            continue
        image_dir.mkdir(exist_ok=True)
        name = f"inline-{index:03d}{suffix}"
        (image_dir / name).write_bytes(blob)
        block["src"] = f"images/{name}"


def source_range(unit: dict[str, Any]) -> list[int]:
    numbers = [int(block.get("sourceParagraph", 0)) for block in unit["elements"] if block.get("sourceParagraph")]
    heading = unit.get("sourceHeadingParagraphs", [])
    numbers.extend(int(value) for value in heading)
    return [min(numbers), max(numbers)] if numbers else [0, 0]


def unit_summary(unit: dict[str, Any]) -> dict[str, Any]:
    keys = ("id", "kind", "label", "title", "shortTitle", "file", "index", "sourceRange", "openingMode")
    return {key: unit[key] for key in keys if key in unit}


def imported_design(reference_path: Path, metadata: dict[str, str], slug: str, cover_source: Path) -> dict[str, Any]:
    raw = json.loads(reference_path.read_text(encoding="utf-8"))
    for key in ("bookSlug", "title", "collection", "series", "coverSpecification", "generationContract", "implementationBoundary"):
        raw.pop(key, None)
    raw["bookSlug"] = slug
    raw["title"] = metadata["title"]
    if metadata.get("collection"):
        raw["collection"] = metadata["collection"]
    if metadata.get("series"):
        raw["series"] = metadata["series"]
    raw["coverSpecification"] = {
        "sourceAsset": {"importedFilename": cover_source.name},
        "imageFitBehavior": "Use object-fit: contain. Preserve complete artwork and aspect ratio; do not crop or stretch.",
    }
    raw["generationContract"] = {
        "contractVersion": IMPORTER_SCHEMA_VERSION,
        "sourceOfTruth": "Imported manuscript normalized into Greyveil source blocks.",
        "sourceFiles": {"bookMetadata": "book.json", "chaptersDirectory": "chapters/", "coverDirectory": "cover/"},
    }
    return raw


def write_json(path: Path, value: dict[str, Any]) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def failed_result(
    job: GenerationJob, manuscript: Path, error: str, slug: str = "", metadata: dict[str, str] | None = None
) -> ImportResult:
    job.fail(error)
    return ImportResult("failed", job, manuscript, slug=slug, metadata=metadata or {}, error=error, warnings=list(job.warnings))


def attention_result(
    job: GenerationJob, manuscript: Path, slug: str, metadata: dict[str, str], warning: str
) -> ImportResult:
    job.add_warning(warning)
    return ImportResult(
        "needs_attention",
        job,
        manuscript,
        slug=slug,
        metadata=metadata,
        warnings=list(job.warnings),
    )
