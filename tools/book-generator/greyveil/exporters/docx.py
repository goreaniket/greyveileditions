"""Print-safe editable DOCX exporter for Greyveil books."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage

from greyveil.exporters.common import (
    ExportTheme,
    block_plain_text,
    book_number_label,
    hex_no_hash,
    hide_folio_for_unit,
    is_quote_block,
    iter_text_runs,
    output_path,
    preferred_cover,
    theme_from_model,
    unit_label,
    unit_phase,
    unit_subtitle,
)
from greyveil.models import BookModel, ChapterBlock, ChapterUnit
from greyveil.utils import first_string, resolve_repo_path


def export_docx(model: BookModel, repo_root: Path, output: Path | None = None) -> Path:
    errors = [issue.message for issue in model.issues if issue.severity == "error"]
    if errors:
        raise ValueError("Cannot export DOCX with validation errors: " + "; ".join(errors))

    destination = output_path(repo_root, "docx", model.slug, "-print-editable.docx", output)
    theme = theme_from_model(model, print_mode=True)

    doc = Document()
    doc.core_properties.title = model.metadata.title
    doc.core_properties.author = model.metadata.author
    doc.core_properties.subject = "Greyveil print/editable edition"
    doc.core_properties.category = "Greyveil Editions"

    configure_page(doc, theme)
    configure_styles(doc, theme)
    add_cover_page(doc, model, theme)
    add_opening_page(doc, model, theme)
    add_units(doc, model, repo_root, theme)

    doc.save(destination)
    return destination


def configure_page(doc: Document, theme: ExportTheme) -> None:
    section = doc.sections[0]
    configure_section_geometry(section, theme)
    enable_mirror_margins(doc)
    configure_section_footer(section, theme, visible=False)


def configure_section_geometry(section, theme: ExportTheme) -> None:
    section.page_width = Inches(theme.trim_width_in)
    section.page_height = Inches(theme.trim_height_in)
    section.left_margin = Inches(theme.inside_margin_in)
    section.right_margin = Inches(theme.outside_margin_in)
    section.top_margin = Inches(theme.top_margin_in)
    section.bottom_margin = Inches(theme.bottom_margin_in)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(theme.folio_bottom_in)
    section.different_first_page_header_footer = False

    set_section_gutter(section, theme.gutter_in)


def set_section_gutter(section, gutter_in: float) -> None:
    try:
        section.gutter = Inches(gutter_in)
        return
    except AttributeError:
        pass

    sect_pr = section._sectPr
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        sect_pr.append(pg_mar)
    pg_mar.set(qn("w:gutter"), str(round(gutter_in * 1440)))


def enable_mirror_margins(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))


def configure_section_footer(section, theme: ExportTheme, *, visible: bool) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_element = footer._element
    for child in list(footer_element):
        footer_element.remove(child)
    footer_element.append(OxmlElement("w:p"))
    if not visible:
        return

    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    set_run_font(run, theme.sans_font, theme.folio_size_pt, theme.subtle)
    add_page_field(run)


def add_folio_section(doc: Document, theme: ExportTheme, *, visible: bool) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_geometry(section, theme)
    configure_section_footer(section, theme, visible=visible)


def add_cover_page(doc: Document, model: BookModel, theme: ExportTheme) -> bool:
    cover = preferred_cover(model.cover_assets, roles=("print", "source", "web"))
    if not cover or not cover.resolved_path or not cover.resolved_path.exists():
        return False

    with PILImage.open(cover.resolved_path) as source:
        image_width, image_height = source.size
    available_width = theme.content_width_in
    available_height = theme.trim_height_in - theme.top_margin_in - theme.bottom_margin_in
    width_scale = available_width / image_width
    height_scale = available_height / image_height
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    if width_scale <= height_scale:
        display_width = available_width
        display_height = image_height * width_scale
        run.add_picture(str(cover.resolved_path), width=Inches(display_width))
    else:
        display_height = available_height
        run.add_picture(str(cover.resolved_path), height=Inches(display_height))
    paragraph.paragraph_format.space_before = Inches(max(0, (available_height - display_height) / 2))
    disable_image_compression(doc)
    doc.add_page_break()
    return True


def disable_image_compression(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:doNotAutoCompressPictures")) is None:
        settings.append(OxmlElement("w:doNotAutoCompressPictures"))


def configure_styles(doc: Document, theme: ExportTheme) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, theme.body_font, theme.body_size_pt, theme.text)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(theme.first_line_indent_in)
    normal.paragraph_format.line_spacing = Pt(theme.body_line_pt)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(theme.paragraph_spacing_pt)

    title = get_or_add_style(styles, "Greyveil Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, theme.display_font, min(theme.title_size_pt, 43), theme.heading, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.line_spacing = 0.94
    title.paragraph_format.space_after = Pt(12)

    subtitle = get_or_add_style(styles, "Greyveil Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle, theme.display_font, 14.2, theme.muted, italic=True)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.line_spacing = 1.18
    subtitle.paragraph_format.space_after = Pt(20)

    kicker = get_or_add_style(styles, "Greyveil Kicker", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(kicker, theme.sans_font, theme.chapter_number_size_pt, theme.accent, bold=True)
    kicker.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kicker.paragraph_format.first_line_indent = Inches(0)
    kicker.paragraph_format.space_after = Pt(theme.opening_kicker_gap_pt)

    heading = get_or_add_style(styles, "Greyveil Unit Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(heading, theme.display_font, theme.chapter_title_size_pt, theme.heading, bold=True)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.first_line_indent = Inches(0)
    heading.paragraph_format.line_spacing = 1.03
    heading.paragraph_format.space_after = Pt(theme.opening_body_gap_pt)

    body = get_or_add_style(styles, "Greyveil Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, theme.body_font, theme.body_size_pt, theme.text)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Inches(theme.first_line_indent_in)
    body.paragraph_format.line_spacing = Pt(theme.body_line_pt)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(theme.paragraph_spacing_pt)

    front = get_or_add_style(styles, "Greyveil Front Matter", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(front, theme.body_font, max(9.8, theme.body_size_pt - 1.0), theme.text)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    front.paragraph_format.first_line_indent = Inches(0)
    front.paragraph_format.line_spacing = Pt(16)
    front.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    front.paragraph_format.space_after = Pt(5)

    contents = get_or_add_style(styles, "Greyveil Contents", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(contents, theme.sans_font, theme.contents_size_pt, theme.text)
    contents.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contents.paragraph_format.first_line_indent = Inches(0)
    contents.paragraph_format.line_spacing = Pt(13)
    contents.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    contents.paragraph_format.space_after = Pt(3)

    quote = get_or_add_style(styles, "Greyveil Quote", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(quote, theme.display_font, 13.2, theme.accent, italic=True)
    quote.paragraph_format.left_indent = Inches(0.22)
    quote.paragraph_format.right_indent = Inches(0.18)
    quote.paragraph_format.first_line_indent = Inches(0)
    quote.paragraph_format.line_spacing = Pt(18)
    quote.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    quote.paragraph_format.space_before = Pt(10)
    quote.paragraph_format.space_after = Pt(10)

    dedication = get_or_add_style(styles, "Greyveil Dedication", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(dedication, theme.display_font, 19.5, theme.accent, italic=True)
    dedication.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication.paragraph_format.first_line_indent = Inches(0)
    dedication.paragraph_format.line_spacing = 1.18
    dedication.paragraph_format.space_after = Pt(10)


def add_opening_page(doc: Document, model: BookModel, theme: ExportTheme) -> None:
    first = doc.add_paragraph()
    first.style = "Greyveil Kicker"
    first.paragraph_format.space_before = Inches(theme.title_opening_top_in)
    if model.metadata.series:
        add_text_run(first, model.metadata.series.upper(), theme.sans_font, theme.chapter_number_size_pt, theme.warm, bold=True)
    if model.metadata.book_number:
        p = doc.add_paragraph(style="Greyveil Kicker")
        add_text_run(p, book_number_label(model.metadata.book_number).upper(), theme.sans_font, theme.chapter_number_size_pt, theme.warm, bold=True)

    add_displaced_rule(
        doc,
        theme,
        width_a=theme.title_rule_primary_in,
        width_b=theme.title_rule_secondary_in,
        left=0.0,
        horizontal_offset=theme.title_rule_horizontal_offset_in,
        vertical_offset=theme.title_rule_vertical_offset_in,
        after_pt=12,
    )

    title = doc.add_paragraph(style="Greyveil Title")
    add_text_run(title, model.metadata.title.upper(), theme.display_font, min(theme.title_size_pt, 43), theme.heading, bold=True)

    if model.metadata.subtitle:
        subtitle = doc.add_paragraph(style="Greyveil Subtitle")
        add_text_run(subtitle, model.metadata.subtitle, theme.display_font, 14.2, theme.muted, italic=True)

    if model.metadata.author:
        author = doc.add_paragraph(style="Greyveil Kicker")
        author.paragraph_format.space_before = Pt(14)
        add_text_run(author, model.metadata.author.upper(), theme.sans_font, 8.5, theme.accent, bold=True)

    if model.metadata.publisher:
        publisher = doc.add_paragraph(style="Greyveil Kicker")
        imprint = model.metadata.publisher.upper()
        if model.metadata.edition_year:
            imprint = f"{imprint} {model.metadata.edition_year}"
        add_text_run(publisher, imprint, theme.sans_font, 8.0, theme.accent, bold=True)


def add_units(doc: Document, model: BookModel, repo_root: Path, theme: ExportTheme) -> None:
    folio_hidden = True
    for unit in model.chapters:
        if unit.kind == "opening":
            continue
        unit_folio_hidden = hide_folio_for_unit(unit, theme)
        if unit_folio_hidden != folio_hidden:
            add_folio_section(doc, theme, visible=not unit_folio_hidden)
        else:
            doc.add_page_break()
        folio_hidden = unit_folio_hidden

        add_unit_header(doc, unit, theme)
        no_indent_next = True
        for block in unit.blocks:
            added_paragraph = add_block(doc, block, unit, model, repo_root, theme, no_indent_next)
            if block.type in {"space", "section-break", "divider"} or is_quote_block(block.type):
                no_indent_next = True
            elif added_paragraph is not None:
                no_indent_next = False


def add_unit_header(doc: Document, unit: ChapterUnit, theme: ExportTheme) -> None:
    if unit.kind == "dedication":
        add_spacer(doc, 1.25)
        add_displaced_rule(
            doc,
            theme,
            width_a=theme.unit_rule_primary_in,
            width_b=theme.unit_rule_secondary_in,
            left=(theme.content_width_in - theme.unit_rule_primary_in) / 2,
            horizontal_offset=theme.unit_rule_horizontal_offset_in,
            vertical_offset=theme.unit_rule_vertical_offset_in,
            after_pt=10,
        )
        title = doc.add_paragraph(style="Greyveil Dedication")
        add_text_run(title, unit.title, theme.display_font, 19.5, theme.accent, italic=True)
        return

    add_spacer(doc, theme.unit_opening_top_in)
    add_displaced_rule(
        doc,
        theme,
        width_a=theme.unit_rule_primary_in,
        width_b=theme.unit_rule_secondary_in,
        left=0.0,
        horizontal_offset=theme.unit_rule_horizontal_offset_in,
        vertical_offset=theme.unit_rule_vertical_offset_in,
        after_pt=theme.opening_kicker_gap_pt,
    )

    phase = unit_phase(unit)
    if phase:
        p = doc.add_paragraph(style="Greyveil Kicker")
        add_text_run(p, phase.upper(), theme.sans_font, theme.chapter_number_size_pt, theme.accent, bold=True)
    elif unit.kind == "chapter":
        p = doc.add_paragraph(style="Greyveil Kicker")
        add_text_run(p, unit_label(unit).upper(), theme.sans_font, theme.chapter_number_size_pt, theme.accent, bold=True)
    elif unit.kind:
        p = doc.add_paragraph(style="Greyveil Kicker")
        add_text_run(p, unit.kind.upper(), theme.sans_font, theme.chapter_number_size_pt, theme.accent, bold=True)

    title = doc.add_paragraph(style="Greyveil Unit Title")
    add_text_run(title, unit.title, theme.display_font, theme.chapter_title_size_pt, theme.heading, bold=True)
    subtitle = unit_subtitle(unit)
    if subtitle:
        p = doc.add_paragraph(style="Greyveil Subtitle")
        add_text_run(p, subtitle, theme.display_font, 14.2, theme.muted, italic=True)


def add_block(
    doc: Document,
    block: ChapterBlock,
    unit: ChapterUnit,
    model: BookModel,
    repo_root: Path,
    theme: ExportTheme,
    no_indent: bool,
):
    if block.type == "space":
        add_spacer(doc, 0.13)
        return None

    if block.type in {"section-break", "divider"}:
        add_displaced_rule(
            doc,
            theme,
            width_a=theme.divider_primary_width_in,
            width_b=theme.divider_secondary_width_in,
            left=0.1,
            horizontal_offset=theme.divider_horizontal_offset_in,
            vertical_offset=theme.divider_vertical_offset_in,
            after_pt=8,
        )
        return None

    if block.type == "image":
        return add_image_block(doc, block, model, repo_root, theme)

    text = block_plain_text(block)
    if not text:
        return None

    style_name = "Greyveil Body"
    if is_quote_block(block.type):
        style_name = "Greyveil Quote"
    elif block.type.startswith("toc-"):
        style_name = "Greyveil Contents"
    elif block.type == "heading":
        style_name = "Greyveil Unit Title"
    elif unit.kind in {"frontmatter", "contents"}:
        style_name = "Greyveil Front Matter"
    elif unit.kind == "dedication":
        style_name = "Greyveil Dedication"

    paragraph = doc.add_paragraph(style=style_name)
    if no_indent or style_name != "Greyveil Body":
        paragraph.paragraph_format.first_line_indent = Inches(0)
    add_block_runs(paragraph, block, theme, style_name)
    return paragraph


def add_image_block(doc: Document, block: ChapterBlock, model: BookModel, repo_root: Path, theme: ExportTheme):
    source = first_string(block.raw.get("src"), block.raw.get("path"), block.raw.get("url"))
    resolved = resolve_repo_path(repo_root, model.root_path, source)
    if not resolved or not resolved.exists():
        return None
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(resolved), width=Inches(theme.content_width_in))
    return paragraph


def add_spacer(doc: Document, height_in: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Inches(height_in)
    paragraph.paragraph_format.line_spacing = 1


def add_displaced_rule(
    doc: Document,
    theme: ExportTheme,
    *,
    width_a: float,
    width_b: float,
    left: float,
    horizontal_offset: float,
    vertical_offset: float,
    after_pt: float,
) -> None:
    add_rule(doc, theme, left=left, width=width_a, color=theme.accent, before_pt=0, after_pt=0)
    add_rule(
        doc,
        theme,
        left=left + horizontal_offset,
        width=width_b,
        color=theme.accent,
        before_pt=vertical_offset * 72,
        after_pt=after_pt,
    )


def add_rule(
    doc: Document,
    theme: ExportTheme,
    *,
    left: float,
    width: float,
    color: str,
    before_pt: float,
    after_pt: float,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(max(0, left))
    paragraph.paragraph_format.right_indent = Inches(max(0, theme.content_width_in - left - width))
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.line_spacing = Pt(1)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), hex_no_hash(color))
    p_bdr.append(top)


def add_block_runs(paragraph, block: ChapterBlock, theme: ExportTheme, style_name: str) -> None:
    font_name, size_pt, color = run_tokens_for_style(style_name, theme)
    for text, bold, italic in iter_text_runs(block):
        run = paragraph.add_run(text)
        set_run_font(
            run,
            font_name,
            size_pt,
            color,
            bold=True if bold else None,
            italic=True if italic else None,
        )


def run_tokens_for_style(style_name: str, theme: ExportTheme) -> tuple[str, float, str]:
    if style_name == "Greyveil Quote":
        return theme.display_font, 13.2, theme.accent
    if style_name == "Greyveil Contents":
        return theme.sans_font, theme.contents_size_pt, theme.text
    if style_name == "Greyveil Unit Title":
        return theme.display_font, theme.chapter_title_size_pt, theme.heading
    if style_name == "Greyveil Front Matter":
        return theme.body_font, max(9.8, theme.body_size_pt - 1.0), theme.text
    if style_name == "Greyveil Dedication":
        return theme.display_font, 19.5, theme.accent
    return theme.body_font, theme.body_size_pt, theme.text


def add_text_run(
    paragraph,
    text: str,
    font_name: str,
    size_pt: float,
    color: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt, color, bold=bold, italic=italic)


def set_run_font(
    run,
    font_name: str,
    size_pt: float,
    color: str,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(hex_no_hash(color))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def set_style_font(style, font_name: str, size_pt: float, color: str, *, bold: bool = False, italic: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor.from_string(hex_no_hash(color))
    style.font.bold = bold
    style.font.italic = italic
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def get_or_add_style(styles, name: str, style_type):
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, style_type)


def add_page_field(run) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)
